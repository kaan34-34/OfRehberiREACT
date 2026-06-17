from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "RAPOR.md"
DOCX_PATH = ROOT / "RAPOR.docx"


def set_cell_text(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def set_style_font(style, size, color=None, bold=False):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
      style.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 15, "2E74B5", 10, 5),
        ("Heading 2", 12.5, "2E74B5", 8, 4),
        ("Heading 3", 11.5, "1F4D78", 6, 3),
    ]:
        style = styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Report Small" not in styles:
        small = styles.add_style("Report Small", WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(small, 9.5, "555555")
        small.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc, lines):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("Of Rehberi Web Uygulaması\nProje Raporu")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(46, 116, 181)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("İleri Web Programlama - Yüksek Lisans")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(31, 77, 120)

    for line in lines[2:8]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(1.0)
        p.paragraph_format.space_after = Pt(8)
        clean = line.replace("**", "")
        parts = clean.split(":", 1)
        if len(parts) == 2:
            set_cell_text(p, parts[0] + ":", bold=True)
            set_cell_text(p, parts[1])
        else:
            set_cell_text(p, clean)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Teslim klasörü: SourceCode/frontend, SourceCode/backend, SourceCode/database")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_page_break()


def add_formatted_paragraph(doc, text, style=None, alignment=None):
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment

    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        if part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(31, 77, 120)
        elif part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="Report Small")
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def build():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover(doc, lines[:8])

    in_code = False
    code_lines = []

    for raw in lines[7:]:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            add_formatted_paragraph(doc, line[2:], "Title", WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith("## "):
            add_formatted_paragraph(doc, line[3:], "Heading 1")
        elif line.startswith("### "):
            add_formatted_paragraph(doc, line[4:], "Heading 2")
        elif line.startswith("- "):
            add_formatted_paragraph(doc, line[2:], "List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_formatted_paragraph(doc, re.sub(r"^\d+\. ", "", line), "List Number")
        else:
            add_formatted_paragraph(doc, line)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
